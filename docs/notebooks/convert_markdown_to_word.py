import os
import re
from pathlib import Path
import markdown
# Fix the import statement to avoid the conflict
import docx
from docx.api import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
import requests
from bs4 import BeautifulSoup
import argparse

def add_hyperlink(paragraph, url, text):
    """
    Add a hyperlink to a paragraph.
    
    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL to link to
        text: The text to display for the hyperlink
    """
    # This gets access to the document's internal relation collection
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)
    
    # Create a run element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    
    # Create a text element
    text_element = docx.oxml.shared.OxmlElement('w:t')
    text_element.text = text
    
    # Add the text element to the run
    new_run.append(text_element)
    
    # Add the run to the hyperlink
    hyperlink.append(new_run)
    
    # Add the hyperlink to the paragraph
    paragraph._p.append(hyperlink)

def process_markdown_file(file_path, doc, image_base_dir):
    """Process a single markdown file and add its content to the document."""
    print(f"Processing {os.path.basename(file_path)}...")
    
    # Read the markdown content
    with open(file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Get the file base name for heading and image path
    file_base_name = os.path.splitext(os.path.basename(file_path))[0]
    
    # Determine image subdirectory based on file name
    img_subdir = os.path.join(os.path.dirname(file_path), f"{file_base_name}_files")
    
    # Add the file name as a heading (unless it's README, then use a more descriptive title)
    if file_base_name.lower() == "readme":
        # For README in root, use "Project Overview", for others use parent folder name
        if os.path.dirname(file_path) == os.path.dirname(os.path.dirname(image_base_dir)):
            doc.add_heading("Project Overview", level=1)
        else:
            parent_folder = os.path.basename(os.path.dirname(file_path))
            doc.add_heading(f"{parent_folder} Documentation", level=1)
    else:
        doc.add_heading(file_base_name, level=1)
    
    # Convert markdown to HTML
    html_content = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    
    # Process HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'code', 'pre', 'ul', 'ol', 'li', 'table', 'img']):
        if element.name.startswith('h'):
            # Get the heading level
            level = int(element.name[1])
            heading = doc.add_heading(element.text, level=level)
            # Adjust font size based on heading level
            if level > 1:  # For headings other than H1, make them smaller
                font_size = 16 - (level * 1)  # Decrease size as level increases
                heading.style.font.size = Pt(font_size)
        
        elif element.name == 'p':
            # Check if paragraph contains links
            links = element.find_all('a')
            if links:
                p = doc.add_paragraph()
                # Split the paragraph by links
                current_pos = 0
                for link in links:
                    # Find position of the link in the original text
                    link_pos = element.text.find(link.text, current_pos)
                    
                    # Add text before the link
                    if link_pos > current_pos:
                        p.add_run(element.text[current_pos:link_pos])
                    
                    # Add the hyperlink
                    add_hyperlink(p, link.get('href', '#'), link.text)
                    
                    # Update current position
                    current_pos = link_pos + len(link.text)
                
                # Add any remaining text after the last link
                if current_pos < len(element.text):
                    p.add_run(element.text[current_pos:])
            else:
                # Regular paragraph without links
                doc.add_paragraph(element.text)
        
        elif element.name == 'pre':
            # Add a code block using a style that exists in the default template
            p = doc.add_paragraph(style='No Spacing')  # Alternative built-in style
            run = p.add_run(element.text)
            run.font.name = 'Courier New'  # Use a monospaced font
            run.font.size = Pt(8)  # Make code font smaller than regular text
        
        elif element.name == 'ul':
            # Unordered list
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        
        elif element.name == 'ol':
            # Ordered list
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
        
        elif element.name == 'img':
            # Handle image references
            img_src = element.get('src', '')
            
            # Look for PNG image format specifically
            img_match = re.search(r'(.+)\.(png|jpg|jpeg|gif)', img_src, re.IGNORECASE)
            if img_match:
                # Extract the base name and extension
                img_base = img_match.group(1)
                img_ext = img_match.group(2)
                
                # Check if the image is in the expected subdirectory
                img_path = os.path.join(img_subdir, f"{os.path.basename(img_base)}.{img_ext}")
                
                # Try alternative path formats if the file doesn't exist
                if not os.path.exists(img_path):
                    # Try finding the image directly referenced in markdown
                    if os.path.exists(os.path.join(os.path.dirname(file_path), img_src)):
                        img_path = os.path.join(os.path.dirname(file_path), img_src)
                    else:
                        print(f"Warning: Image not found at {img_path}, looking for alternatives...")
                        
                        # Look for the image in the files directory with number suffix
                        if os.path.exists(img_subdir):
                            potential_files = [f for f in os.listdir(img_subdir) 
                                            if f.endswith(f'.{img_ext}') and os.path.basename(img_base) in f]
                            
                            if potential_files:
                                img_path = os.path.join(img_subdir, potential_files[0])
                                print(f"Found alternative image: {img_path}")
                            else:
                                print(f"Warning: Could not find image for {img_src}")
                                continue
                        else:
                            print(f"Warning: Image subdirectory {img_subdir} does not exist")
                            continue
                
                # Add the image to the document if it exists
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(6))
                    # Add caption with image alt text if available
                    alt_text = element.get('alt', '')
                    caption_paragraph = doc.add_paragraph(f"Figure: {os.path.basename(img_path)}")
                    caption_paragraph.style = doc.styles['Caption']
                    caption_paragraph.style.font.size = Pt(9)  # Smaller font for captions
                else:
                    print(f"Warning: Image not found at {img_path}")

def convert_markdown_to_word(markdown_dir, output_file, image_base_dir=None):
    """
    Convert multiple markdown files to a single Word document.
    
    Args:
        markdown_dir (str): Directory containing markdown files
        output_file (str): Output Word document path
        image_base_dir (str, optional): Base directory for images, 
                                      defaults to same as markdown files
    """
    if image_base_dir is None:
        image_base_dir = markdown_dir
    
    # Create a new Word document
    doc = Document()
    
    # Customize styles for font sizes
    # Make Heading 1 larger
    style = doc.styles['Heading 1']
    style.font.size = Pt(18)  # Larger font for Heading 1
    style.font.bold = True
    
    # Make normal text smaller
    style = doc.styles['Normal']
    style.font.size = Pt(10)  # Smaller font for regular text
    
    # Start with READMEs at different levels
    markdown_dir_abs = os.path.abspath(markdown_dir)

    # First, try to include the root README (two levels up)
    root_readme_path = os.path.join(os.path.dirname(os.path.dirname(markdown_dir_abs)), "README.md")
    if os.path.exists(root_readme_path):
        print(f"Including root README from: {root_readme_path}")
        process_markdown_file(root_readme_path, doc, image_base_dir)
        doc.add_page_break()
    else:
        print("Root README not found at:", root_readme_path)
    
    # Next, look for a README in the current folder
    folder_readme_path = os.path.join(markdown_dir, "README.md")
    if os.path.exists(folder_readme_path):
        print(f"Including folder README from: {folder_readme_path}")
        process_markdown_file(folder_readme_path, doc, image_base_dir)
        doc.add_page_break()
    else:
        print("Folder README not found at:", folder_readme_path)
    
    # Get all markdown files in the directory (excluding READMEs which we've already processed)
    markdown_files = sorted([f for f in os.listdir(markdown_dir) 
                            if f.endswith('.md') and f != "README.md"])
    
    print(f"Found {len(markdown_files)} markdown files to convert")
    
    # Process each markdown file
    for md_file in markdown_files:
        md_path = os.path.join(markdown_dir, md_file)
        process_markdown_file(md_path, doc, image_base_dir)
        
        # Add a page break after each file except the last one
        if md_file != markdown_files[-1]:
            doc.add_page_break()
    
    # Save the Word document
    doc.save(output_file)
    print(f"Conversion complete! Word document saved as {output_file}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Convert markdown files to Word document')
    parser.add_argument('--markdown_dir', required=True, help='Directory containing markdown files')
    parser.add_argument('--output_file', required=True, help='Output Word document path')
    parser.add_argument('--image_base_dir', help='Base directory for images (optional)')
    
    args = parser.parse_args()
    
    convert_markdown_to_word(
        args.markdown_dir, 
        args.output_file,
        args.image_base_dir
    )